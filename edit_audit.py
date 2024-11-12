from flask import Flask, request, send_file
import io
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

app = Flask(__name__)


def add_bookmark(paragraph, bookmark_name):
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(id(paragraph)))
    bookmark_start.set(qn('w:name'), bookmark_name)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(id(paragraph)))

    paragraph._p.append(bookmark_start)
    paragraph._p.append(bookmark_end)


def add_hyperlink(paragraph, bookmark_name, text):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def replace_image_references(document, image_map):
    for i, paragraph in enumerate(document.paragraphs):
        for key, value in image_map.items():
            search_text = f'(см. {key})'
            replacement_text = f'(см. {value})'

            if search_text in paragraph.text:
                parts = paragraph.text.split(search_text)

                p = paragraph._p
                if p is not None:
                    p.getparent().remove(p)

                next_paragraph = document.paragraphs[i] if i < len(document.paragraphs) - 1 else None

                if next_paragraph is not None:
                    new_paragraph = next_paragraph.insert_paragraph_before()
                else:
                    new_paragraph = document.add_paragraph()  # Вставка в конец, если последний абзац

                add_bookmark(new_paragraph, key)

                new_paragraph.add_run(parts[0])

                add_hyperlink(new_paragraph, key, replacement_text)

                if len(parts) > 1:
                    new_paragraph.add_run(parts[1])

                new_paragraph.style = document.styles['Основной текст icl']


def find_images_and_captions(document):
    images = {}

    for paragraph in document.paragraphs:
        if 'Рисунок' in paragraph.text:
            caption = paragraph.text.split('-')[0].strip()
            name = paragraph.text.split('-')[1].strip()
            images[name] = caption

    return images


@app.route('/edit_docx', methods=['POST'])
def edit_docx():
    file = request.files['file']
    file_content = file.read()
    open('file_from_front.docx', mode='wb').write(file_content)
    document = Document(io.BytesIO(file_content))
    image_counter = 1
    in_section_6 = False
    need_to_delete_next = False
    next_p_add_art = False
    list_paragraph_to_check = ["Описание:", "Наблюдения:", "Ссылки:", "Доказательства:", "Затронутые активы:", "Рекомендации:"]

    for paragraph in document.paragraphs:
        if paragraph.style.name == 'Normal':
            paragraph.style = document.styles['Основной текст icl']

        if paragraph.text.find("оценивается как") != -1:
            start = paragraph.text.find("оценивается как")
            picture = paragraph.text[start+16:].split(",")[0]

        if next_p_add_art:
            run = paragraph.insert_paragraph_before()
            run.alignment = 1
            run.add_run().add_picture(f"arts/{picture}.png", width=Inches(6))
            next_p_add_art = False

        if paragraph.text.find("От внешнего нарушителя ресурса") != -1:
            next_p_add_art = True

        if "Детальное описание хода работ и результатов" in paragraph.text and paragraph.style.name == "Heading 1":
            in_section_6 = True

        if in_section_6:
            if paragraph.text == "" and reserv_paragraph.text in list_paragraph_to_check:
                p = reserv_paragraph._element
                p.getparent().remove(p)
                p._element = None
                p_ = paragraph._element
                p_.getparent().remove(p_)
                p_._element = None

            if need_to_delete_next:
                p = paragraph._element
                p.getparent().remove(p)
                p._element = None
                need_to_delete_next = False

            if paragraph.text == "CVSS балл:  ()":
                p = paragraph._element
                p.getparent().remove(p)
                p._element = None
                need_to_delete_next = True

            if paragraph.style.name == 'Caption':
                paragraph.text = f"Рисунок 6.{image_counter} - {paragraph.text}"
                image_counter += 1

            reserv_paragraph = paragraph

    images = find_images_and_captions(document)
    replace_image_references(document, images)

    # Работа с определениями
    table = document.tables[0]

    for i, column in enumerate(table.columns):
        if i == 0:
            column.width = Cm(2)  # Первый столбец - 2 см
        elif i == 1:
            column.width = Cm(0.5)  # Второй столбец - 0.5 см
        elif i == 2:
            column.width = Cm(13)  # Третий столбец - 13 см

    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(13)

    table.style = 'pwndoc-table'

    for row in table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(nsdecls('w'))))

    # Работа с терминами
    table2 = document.tables[1]

    for i, column in enumerate(table2.columns):
        if i == 0:
            column.width = Cm(4)  # Первый столбец - 2 см
        elif i == 1:
            column.width = Cm(0.5)  # Второй столбец - 0.5 см
        elif i == 2:
            column.width = Cm(11)  # Третий столбец - 13 см

    for row in table2.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(11)

    table2.style = 'pwndoc-table'

    for row in table2.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(nsdecls('w'))))

    processed_file = io.BytesIO()
    document.save(processed_file)
    processed_file.seek(0)

    return send_file(
        processed_file,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='edited_document.docx'
    )


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
