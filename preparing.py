import io
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Cm, Pt
from lxml import etree
from utils import find_images_and_captions, replace_image_references, \
	create_bookmarks, highlight_clear


class NAMESPACES:
	DOCX = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


img_names = ['высокий', 'выше среднего', 'крайне низкий', 'ниже среднего', 'низкий', 'средний']


def change_color_for_bold(document: Document):
	for paragraph in document.paragraphs:
		for run in paragraph.runs:
			for bold in run._element.findall('.//w:b', NAMESPACES.DOCX):
				bold_color = etree.Element("{" + NAMESPACES.DOCX['w'] + "}color", nsmap=NAMESPACES.DOCX)
				bold_color.attrib["{" + NAMESPACES.DOCX['w'] + '}themeColor'] = 'accent1'
				bold_color.attrib["{" + NAMESPACES.DOCX['w'] + '}val'] = '002854'
				bold.getparent().append(bold_color)


def document_preparing(document: Document):
	image_counter = 1
	in_section_6 = False
	need_to_delete_next = False
	next_p_add_art = False
	list_paragraph_to_check = ["Описание:", "Наблюдения:", "Ссылки:", "Доказательства:", "Затронутые активы:",
							   "Рекомендации:"]

	for i, paragraph in enumerate(document.paragraphs):
		if paragraph.style.name == 'Normal':
			paragraph.style = document.styles['Основной текст icl']

		if paragraph.text.find("оценивается как") != -1:
			start = paragraph.text.find("оценивается как")
			picture = paragraph.text[start + 16:].split(",")[0].split('.')[0]
			children = document.paragraphs[i+1]._element.getchildren()
			# document.paragraphs[i + 1]._element.getparent().remove(document.paragraphs[i + 1]._element)
			# run = etree.SubElement(paragraph._element, "{" + NAMESPACES.DOCX['w'] + "}r", nsmap=NAMESPACES.DOCX)
			# t = etree.SubElement(run, "{" + NAMESPACES.DOCX['w'] + "}t", nsmap=NAMESPACES.DOCX)
			# t.text = added_text
			for ch in children:
				paragraph._element.append(ch)

		if next_p_add_art:
			run = paragraph.insert_paragraph_before()
			run.alignment = 1
			run.add_run().add_picture(f"arts/{picture}.png", width=Cm(17.5))
			next_p_add_art = False

		if paragraph.text.find("От внешнего нарушителя ресурса") != -1:
			next_p_add_art = True

		if "Детальное описание хода работ и результатов" in paragraph.text and paragraph.style.name == "Heading 1":
			in_section_6 = True

		if in_section_6:
			if paragraph.text == "" and reserv_paragraph.text in list_paragraph_to_check:
				p = reserv_paragraph._element
				p.getparent().remove(p)
				p._element = None
				p_ = paragraph._element
				p_.getparent().remove(p_)
				p_._element = None

			if need_to_delete_next:
				p = paragraph._element
				p.getparent().remove(p)
				p._element = None
				need_to_delete_next = False

			if paragraph.text == "CVSS балл:  ()":
				p = paragraph._element
				p.getparent().remove(p)
				p._element = None
				need_to_delete_next = True

			if paragraph.style.name == 'Caption':
				text = paragraph.text
				start_b, end_b = create_bookmarks(text)

				for ch in paragraph._element.getchildren():
					if ch.tag.replace('{'+NAMESPACES.DOCX['w']+'}', '') == 'r':
						paragraph._element.remove(ch)
				run = paragraph.add_run()
				run.text = f"Рисунок 6.{image_counter} - {text}"
				run.font.size = Pt(10)

				paragraph._element.append(start_b)
				paragraph._element.append(run._element)
				paragraph._element.append(end_b)

				image_counter += 1

			reserv_paragraph = paragraph

	images = find_images_and_captions(document)
	replace_image_references(document, images)

	# Работа с определениями
	table = document.tables[0]

	for i, column in enumerate(table.columns):
		if i == 0:
			column.width = Cm(2)  # Первый столбец - 2 см
		elif i == 1:
			column.width = Cm(0.5)  # Второй столбец - 0.5 см
		elif i == 2:
			column.width = Cm(13)  # Третий столбец - 13 см

	for row in table.rows:
		row.cells[0].width = Cm(2)
		row.cells[1].width = Cm(0.5)
		row.cells[2].width = Cm(13)

	table.style = 'pwndoc-table'

	for row in table.rows:
		for i, cell in enumerate(row.cells):
			cell._element.get_or_add_tcPr().append(parse_xml(
				r'<w:tcBorders {}><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(
					nsdecls('w'))))
			for p in cell.paragraphs:
				for r in p.runs:
					r.font.name = 'Arial'
				if i == 2:
					p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
			cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


	# Работа с терминами
	table2 = document.tables[1]

	for i, column in enumerate(table2.columns):
		if i == 0:
			column.width = Cm(4)  # Первый столбец - 2 см
		elif i == 1:
			column.width = Cm(0.5)  # Второй столбец - 0.5 см
		elif i == 2:
			column.width = Cm(11)  # Третий столбец - 13 см

	for row in table2.rows:
		row.cells[0].width = Cm(4)
		row.cells[1].width = Cm(0.5)
		row.cells[2].width = Cm(11)

	table2.style = 'pwndoc-table'

	for row in table2.rows:
		for i, cell in enumerate(row.cells):
			cell._element.get_or_add_tcPr().append(parse_xml(
				r'<w:tcBorders {}><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(
					nsdecls('w'))))
			for p in cell.paragraphs:
				for r in p.runs:
					r.font.name = 'Arial'
				if i == 2:
					p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

			cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


	# проверка на Null в полях "Уровень трудности устранения" и "Приоритет" (если Null, то не выводить название)
	# замена пунктиров на точки в рекомендациях
	fields = ['Уровень трудности устранения', 'Приоритет']
	in_section_6 = False
	for p in document._element.findall('.//w:p', NAMESPACES.DOCX):
		text = p.text
		# if text is None:
		# 	text = p.xpath("string()")
		# if text is None:
		# 	text = p.tail
		for field in fields:
			if field in text:
				parts = text.split(': ')
				if len(parts) < 2:
					p.getparent().remove(p)
				elif parts[1].replace(' ', '') == '':
					p.getparent().remove(p)

		if text.lower().replace(' ', '') == 'Детальное описание хода работ и результатов'.lower().replace(' ', ''):
			# проверка на 6 пункт
			in_section_6 = True

		if in_section_6:

			if text.lower().replace(' ', '') == 'Рекомендации:'.lower():
				runs = p.findall('.//w:r', NAMESPACES.DOCX)
				if len(runs) > 0:
					runs[0].text = 'Рекомендации к устранению:'

			p_pr = p.find('.//w:pPr', NAMESPACES.DOCX)
			if p_pr is not None:
				num_id = p_pr.find('.//w:numId', NAMESPACES.DOCX)
				if num_id is not None:
					num_id.attrib["{" + NAMESPACES.DOCX['w'] + '}val'] = "3" # применение нужного стиля пунктира
					p_style = p_pr.find('.//w:pStyle', NAMESPACES.DOCX)
					if p_style is not None:
						p_style.attrib["{" + NAMESPACES.DOCX['w'] + '}val'] = "icl" # применение нужного стиля списка

					# регулировка отступа слева
					ind = etree.Element("{" + NAMESPACES.DOCX['w'] + "}ind", nsmap=NAMESPACES.DOCX)
					ind.attrib["{" + NAMESPACES.DOCX['w'] + '}left'] = "1134" #"720"

					p_style.addnext(ind)

	# удаляем выделения текста
	highlight_clear(document)

	# меняем цвет жирного шрифта на синий
	change_color_for_bold(document)
